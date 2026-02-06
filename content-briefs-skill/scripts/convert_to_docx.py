#!/usr/bin/env python3
"""
Convert Markdown briefs to Word documents (.docx)
Preserves formatting, headings, lists, tables, and code blocks
"""

import sys
import re
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    """
    # Get the run object
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add blue color and underline
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def parse_markdown_to_docx(markdown_text, output_path):
    """
    Convert markdown text to a formatted Word document.
    """
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    in_list = False
    list_level = 0

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_content = []
            else:
                # End code block - add to document
                in_code_block = False
                if code_block_content:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    run = p.add_run('\n'.join(code_block_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                code_block_content = []
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Skip horizontal rules and empty lines at start
        if line.strip() in ['---', '***', '___'] or (not line.strip() and i == 0):
            i += 1
            continue

        # H1 - Main title
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(46, 125, 50)  # Green
            p.runs[0].font.size = Pt(24)
            p.runs[0].font.bold = True

        # H2 - Major sections
        elif line.startswith('## '):
            text = line[3:].strip()
            # Remove emojis from headings
            text = re.sub(r'[^\w\s\-:().,!?&/]', '', text)
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = RGBColor(46, 125, 50)
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.bold = True

        # H3 - Subsections
        elif line.startswith('### '):
            text = line[4:].strip()
            text = re.sub(r'[^\w\s\-:().,!?&/]', '', text)
            p = doc.add_heading(text, level=3)
            p.runs[0].font.color.rgb = RGBColor(51, 51, 51)
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True

        # H4 - Minor subsections
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            p.runs[0].font.size = Pt(12)

        # H5
        elif line.startswith('##### '):
            text = line[6:].strip()
            p = doc.add_heading(text, level=5)
            p.runs[0].font.size = Pt(11)

        # Bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Remove markdown bold/italic
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)

            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)

        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)

            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)

        # Block quotes
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            # Skip empty blockquotes (used for paragraph breaks in multi-line quotes)
            if not text:
                i += 1
                continue
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.runs[0]
            run.italic = True
            run.font.color.rgb = RGBColor(102, 102, 102)

        # Tables (simple markdown tables)
        elif '|' in line and line.strip().startswith('|'):
            # Collect table rows
            table_rows = []
            while i < len(lines) and '|' in lines[i]:
                row = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                # Skip separator rows
                if not all(re.match(r'^[-:]+$', cell.strip()) for cell in row):
                    table_rows.append(row)
                i += 1

            if table_rows:
                # Create table
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Light Grid Accent 1'

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data

                        # Header row
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True

                doc.add_paragraph()  # Add space after table
            continue

        # Regular paragraphs
        elif line.strip():
            text = line.strip()

            # Handle inline formatting
            # Bold
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            p = doc.add_paragraph()

            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    # Check for inline code
                    code_parts = re.split(r'(`[^`]+`)', part)
                    for code_part in code_parts:
                        if code_part.startswith('`') and code_part.endswith('`'):
                            run = p.add_run(code_part[1:-1])
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                        elif code_part:
                            run = p.add_run(code_part)

            p.paragraph_format.space_after = Pt(6)

        # Empty line
        else:
            doc.add_paragraph()

        i += 1

    # Save document
    doc.save(output_path)
    print(f"[OK] Created: {output_path}")


def convert_file(input_path):
    """
    Convert a single markdown file to docx.
    Returns tuple of (success: bool, input_path: Path)
    """
    input_path = Path(input_path)

    if not input_path.exists():
        print(f"[ERROR] File not found: {input_path}")
        return False, input_path

    # Read markdown content
    with open(input_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    # Generate output path
    output_path = input_path.with_suffix('.docx')

    # Convert
    try:
        parse_markdown_to_docx(markdown_content, str(output_path))
        return True, input_path
    except Exception as e:
        print(f"[ERROR] Error converting {input_path}: {e}")
        return False, input_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown briefs to Word documents',
        usage='%(prog)s [files...] [--all] [--cleanup-md]'
    )
    parser.add_argument('files', nargs='*', help='Markdown files to convert')
    parser.add_argument('--all', action='store_true', help='Convert all .md files in output folder')
    parser.add_argument('--cleanup-md', action='store_true', help='Remove source .md files after successful conversion')

    args = parser.parse_args()

    # Validate arguments
    if not args.files and not args.all:
        parser.print_help()
        sys.exit(1)

    files_to_convert = []
    cleanup_md = args.cleanup_md
    cleaned_files = []

    if args.all:
        # Convert all markdown files in output folder
        script_dir = Path(__file__).parent
        output_dir = script_dir.parent / 'output'

        files_to_convert = list(output_dir.glob('*.md'))

        if not files_to_convert:
            print("No markdown files found in output folder.")
            sys.exit(0)

        print(f"Converting {len(files_to_convert)} markdown files to Word...\n")
    else:
        # Convert specified files
        files_to_convert = args.files
        if len(files_to_convert) > 1:
            print(f"Converting {len(files_to_convert)} markdown files to Word...\n")

    success_count = 0
    for md_file in files_to_convert:
        success, input_path = convert_file(md_file)

        if success:
            success_count += 1

            # Clean up .md file if requested and conversion was successful
            if cleanup_md:
                try:
                    input_path.unlink()
                    cleaned_files.append(str(input_path))
                    print(f"[CLEANUP] Removed: {input_path}")
                except Exception as e:
                    print(f"[WARNING] Could not delete {input_path}: {e}")

    # Print summary
    if args.all or len(files_to_convert) > 1:
        print(f"\n[SUCCESS] Converted {success_count}/{len(files_to_convert)} files")

    if cleaned_files:
        print(f"[CLEANUP] Removed {len(cleaned_files)} markdown file(s)")
        for cleaned in cleaned_files:
            print(f"  - {cleaned}")


if __name__ == '__main__':
    main()
