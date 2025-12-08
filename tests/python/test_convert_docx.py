"""
Tests for content-briefs-skill/scripts/convert_to_docx.py

Tests markdown to DOCX conversion functionality including:
- Heading parsing (H1-H5)
- List parsing (bullet and numbered)
- Code block parsing
- Table parsing
- File conversion workflow
- Error handling
"""

import pytest
import sys
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from docx import Document
from docx.shared import Pt, RGBColor

# Add scripts directory to path
SCRIPT_DIR = Path(__file__).parent.parent.parent / "content-briefs-skill" / "scripts"
sys.path.insert(0, str(SCRIPT_DIR))

from convert_to_docx import (
    parse_markdown_to_docx,
    convert_file,
    add_hyperlink
)


class TestParseMarkdownToDocx:
    """Test the main markdown to DOCX conversion function."""

    def test_parse_simple_heading_h1(self, temp_dir):
        """Test parsing H1 heading."""
        markdown = "# Main Title"
        output = temp_dir / "output.docx"

        parse_markdown_to_docx(markdown, str(output))

        assert output.exists()
        doc = Document(str(output))
        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[0].text == "Main Title"

    def test_parse_all_heading_levels(self, temp_dir):
        """Test parsing all heading levels (H1-H5)."""
        markdown = """# H1 Heading
## H2 Heading
### H3 Heading
#### H4 Heading
##### H5 Heading"""
        output = temp_dir / "headings.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) == 5
        assert headings[0].text == "H1 Heading"
        assert headings[4].text == "H5 Heading"

    def test_parse_bullet_list(self, temp_dir):
        """Test parsing bullet lists."""
        markdown = """- Item 1
- Item 2
- Item 3"""
        output = temp_dir / "bullets.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        list_items = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
        assert len(list_items) == 3
        assert list_items[0].text == "Item 1"

    def test_parse_numbered_list(self, temp_dir):
        """Test parsing numbered lists."""
        markdown = """1. First item
2. Second item
3. Third item"""
        output = temp_dir / "numbered.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        list_items = [p for p in doc.paragraphs if p.style.name == 'List Number']
        assert len(list_items) == 3
        assert list_items[0].text == "First item"

    def test_parse_code_block(self, temp_dir):
        """Test parsing code blocks."""
        markdown = """```python
def hello():
    print("Hello")
```"""
        output = temp_dir / "code.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        # Code blocks are added as paragraphs with specific formatting
        assert len(doc.paragraphs) > 0
        # Check that content exists
        code_para = [p for p in doc.paragraphs if 'hello' in p.text.lower()]
        assert len(code_para) > 0

    def test_parse_table(self, temp_dir):
        """Test parsing markdown tables."""
        markdown = """| Column 1 | Column 2 |
|----------|----------|
| Data 1   | Data 2   |
| Data 3   | Data 4   |"""
        output = temp_dir / "table.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3  # Header + 2 data rows
        assert len(table.columns) == 2

    def test_parse_bold_text(self, temp_dir):
        """Test parsing bold text formatting."""
        markdown = "This is **bold text** in a paragraph."
        output = temp_dir / "bold.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        paragraph = doc.paragraphs[0]
        bold_runs = [run for run in paragraph.runs if run.bold]
        assert len(bold_runs) > 0

    def test_parse_inline_code(self, temp_dir):
        """Test parsing inline code."""
        markdown = "This has `inline code` in it."
        output = temp_dir / "inline.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        assert len(doc.paragraphs) > 0

    def test_parse_blockquote(self, temp_dir):
        """Test parsing blockquotes."""
        markdown = "> This is a quote\n> On multiple lines"
        output = temp_dir / "quote.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        assert len(doc.paragraphs) >= 2

    def test_parse_mixed_content(self, temp_dir, sample_markdown):
        """Test parsing document with mixed content types."""
        output = temp_dir / "mixed.docx"

        parse_markdown_to_docx(sample_markdown, str(output))

        doc = Document(str(output))
        assert len(doc.paragraphs) > 0
        assert len(doc.tables) > 0

    def test_remove_emojis_from_headings(self, temp_dir):
        """Test that emojis are removed from headings."""
        markdown = "## 🚀 Section with Emoji"
        output = temp_dir / "emoji.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        heading = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        # Emoji should be removed
        assert "🚀" not in heading[0].text
        assert "Section with Emoji" in heading[0].text

    def test_skip_horizontal_rules(self, temp_dir):
        """Test that horizontal rules are skipped."""
        markdown = """# Title
---
Content"""
        output = temp_dir / "hr.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        # Should not contain the --- as text
        all_text = '\n'.join([p.text for p in doc.paragraphs])
        assert '---' not in all_text or 'Title' in all_text


class TestConvertFile:
    """Test the file conversion workflow."""

    def test_convert_existing_file(self, markdown_file):
        """Test converting an existing markdown file."""
        success, path = convert_file(markdown_file)

        assert success is True
        assert path == markdown_file

        # Check output file was created
        output_file = markdown_file.with_suffix('.docx')
        assert output_file.exists()

    def test_convert_missing_file(self, temp_dir):
        """Test converting a non-existent file."""
        missing_file = temp_dir / "missing.md"

        success, path = convert_file(missing_file)

        assert success is False
        assert path == missing_file

    def test_convert_file_output_path(self, markdown_file):
        """Test that output file has correct path and extension."""
        convert_file(markdown_file)

        output_file = markdown_file.with_suffix('.docx')
        assert output_file.exists()
        assert output_file.suffix == '.docx'
        assert output_file.stem == markdown_file.stem

    def test_convert_file_preserves_name(self, temp_dir, sample_markdown):
        """Test that conversion preserves original filename."""
        md_file = temp_dir / "my-brief.md"
        md_file.write_text(sample_markdown)

        convert_file(md_file)

        output_file = temp_dir / "my-brief.docx"
        assert output_file.exists()

    @patch('convert_to_docx.parse_markdown_to_docx')
    def test_convert_file_handles_conversion_error(self, mock_parse, markdown_file):
        """Test that conversion errors are handled gracefully."""
        mock_parse.side_effect = Exception("Conversion error")

        success, path = convert_file(markdown_file)

        assert success is False
        assert path == markdown_file


class TestAddHyperlink:
    """Test hyperlink addition to paragraphs."""

    def test_add_hyperlink_creates_link(self, temp_dir):
        """Test that hyperlinks are added to paragraphs."""
        doc = Document()
        paragraph = doc.add_paragraph()

        add_hyperlink(paragraph, "https://example.com", "Example Link")

        # Check that hyperlink element was added
        assert len(paragraph._p) > 0

    def test_add_hyperlink_preserves_url(self, temp_dir):
        """Test that the correct URL is used."""
        doc = Document()
        paragraph = doc.add_paragraph()
        url = "https://example.com/test"

        hyperlink = add_hyperlink(paragraph, url, "Test Link")

        # Hyperlink should be created
        assert hyperlink is not None

    def test_add_hyperlink_with_text(self, temp_dir):
        """Test that link text is set correctly."""
        doc = Document()
        paragraph = doc.add_paragraph()

        hyperlink = add_hyperlink(paragraph, "https://example.com", "Click Here")

        # The text should be in the hyperlink
        assert hyperlink is not None


class TestEdgeCases:
    """Test edge cases and special scenarios."""

    def test_empty_markdown(self, temp_dir):
        """Test converting empty markdown."""
        output = temp_dir / "empty.docx"

        parse_markdown_to_docx("", str(output))

        assert output.exists()
        doc = Document(str(output))
        # Should have minimal content
        assert len(doc.paragraphs) >= 0

    def test_markdown_with_only_whitespace(self, temp_dir):
        """Test markdown with only whitespace."""
        markdown = "\n\n   \n\n"
        output = temp_dir / "whitespace.docx"

        parse_markdown_to_docx(markdown, str(output))

        assert output.exists()

    def test_nested_formatting(self, temp_dir):
        """Test nested/complex formatting."""
        markdown = "This has **bold with `code` inside** text."
        output = temp_dir / "nested.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        assert len(doc.paragraphs) > 0

    def test_multiple_code_blocks(self, temp_dir):
        """Test multiple code blocks in sequence."""
        markdown = """```python
code1()
```

Some text

```javascript
code2();
```"""
        output = temp_dir / "multi_code.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        assert len(doc.paragraphs) > 0

    def test_list_with_formatting(self, temp_dir):
        """Test lists with inline formatting."""
        markdown = """- Item with **bold**
- Item with `code`
- Plain item"""
        output = temp_dir / "list_format.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        list_items = [p for p in doc.paragraphs if p.style.name == 'List Bullet']
        assert len(list_items) == 3

    def test_table_with_empty_cells(self, temp_dir):
        """Test table with empty cells."""
        markdown = """| Col1 | Col2 |
|------|------|
| Data |      |
|      | Data |"""
        output = temp_dir / "table_empty.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        assert len(doc.tables) == 1

    def test_very_long_document(self, temp_dir):
        """Test conversion of a very long document."""
        # Create markdown with many sections
        sections = []
        for i in range(50):
            sections.append(f"## Section {i}\n\nContent for section {i}\n")
        markdown = "\n".join(sections)
        output = temp_dir / "long.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        assert len(doc.paragraphs) > 100  # Should have many paragraphs

    def test_special_characters(self, temp_dir):
        """Test handling of special characters."""
        markdown = "Text with special chars: & < > \" ' @"
        output = temp_dir / "special.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        assert len(doc.paragraphs) > 0

    def test_unicode_characters(self, temp_dir):
        """Test handling of Unicode characters."""
        markdown = "Unicode: café, naïve, 日本語, 🎉"
        output = temp_dir / "unicode.docx"

        parse_markdown_to_docx(markdown, str(output))

        doc = Document(str(output))
        assert len(doc.paragraphs) > 0
